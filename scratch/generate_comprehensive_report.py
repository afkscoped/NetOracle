import os
import sys
import re
import html
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether, Preformatted
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# Define paths
MD_PATH = Path(r"C:\Users\Rishab Nayak\.gemini\antigravity\brain\82c3609a-f217-4264-abce-7ee71b3d74b9\NetOracle_Complete_Technical_Report.md")
REPORTS_DIR = Path("reports")
REPORTS_DIR.mkdir(parents=True, exist_ok=True)
DOCX_OUT_PATH = REPORTS_DIR / "NetOracle_Complete_Technical_Report.docx"
PDF_OUT_PATH = REPORTS_DIR / "NetOracle_Complete_Technical_Report.pdf"

ARTIFACT_DIR = Path(r"C:\Users\Rishab Nayak\.gemini\antigravity\brain\82c3609a-f217-4264-abce-7ee71b3d74b9")
DOCX_ART_PATH = ARTIFACT_DIR / "NetOracle_Complete_Technical_Report.docx"
PDF_ART_PATH = ARTIFACT_DIR / "NetOracle_Complete_Technical_Report.pdf"

# Register Arial from Windows system fonts directory for PDF
try:
    pdfmetrics.registerFont(TTFont('Arial', 'C:\\Windows\\Fonts\\arial.ttf'))
    pdfmetrics.registerFont(TTFont('Arial-Bold', 'C:\\Windows\\Fonts\\arialbd.ttf'))
    pdfmetrics.registerFont(TTFont('Arial-Italic', 'C:\\Windows\\Fonts\\ariali.ttf'))
    pdfmetrics.registerFont(TTFont('Arial-BoldItalic', 'C:\\Windows\\Fonts\\arialbi.ttf'))
    pdfmetrics.registerFont(TTFont('Consolas', 'C:\\Windows\\Fonts\\consola.ttf'))
    pdfmetrics.registerFont(TTFont('Consolas-Bold', 'C:\\Windows\\Fonts\\consolab.ttf'))
    
    FONT_NORMAL = 'Arial'
    FONT_BOLD = 'Arial-Bold'
    FONT_ITALIC = 'Arial-Italic'
    FONT_MONO = 'Consolas'
    print("Successfully registered system TTF fonts for PDF.")
except Exception as e:
    print(f"Failed to register system TTF fonts: {e}. Falling back to standard Helvetica/Courier.")
    FONT_NORMAL = 'Helvetica'
    FONT_BOLD = 'Helvetica-Bold'
    FONT_ITALIC = 'Helvetica-Oblique'
    FONT_MONO = 'Courier'

# Special character replacement for PDF rendering stability
def replace_special_characters(text):
    replacements = {
        '⊙': '*',
        '◦': '*',
        'ℝ': 'R',
        'ᵀ': '^T',
        'ᵢ': '_i',
        'ₙ': '_n',
        'ᶻ': '^z',
        'ᵉ': '^e',
        '½': '1/2',
        '⁻': '-',
        'ᶠ': '^f',
        '→': '->',
        '←': '<-',
        '≤': '<=',
        '≥': '>=',
        '≈': '~',
        '≠': '!=',
        '•': '*',
        '┌': '+',
        '─': '-',
        '┐': '+',
        '│': '|',
        '├': '+',
        '┬': '+',
        '┬': '+',
        '┤': '|',
        '└': '+',
        '┴': '+',
        '┘': '+',
        '┼': '+',
        '▲': '^',
        '▼': 'v',
        '❖': '*',
        'Σ': 'Sum',
        'Φ': 'Phi',
        'α': 'alpha',
        'β': 'beta',
        'γ': 'gamma',
        'δ': 'delta',
        'ε': 'epsilon',
        'θ': 'theta',
        'λ': 'lambda',
        'μ': 'mu',
        'π': 'pi',
        'σ': 'sigma',
        'ω': 'omega',
        '∂': 'd',
        'ŷ': 'y_hat',
        '₁': '_1',
        '₂': '_2',
        '₃': '_3',
        '₄': '_4',
        '₅': '_5',
        'ₘ': '_m',
        'ₜ': '_t',
        '₀': '_0'
    }
    for orig, rep in replacements.items():
        text = text.replace(orig, rep)
    return text

def html_escape_and_md_to_html(text):
    text = replace_special_characters(text)
    escaped = html.escape(text)
    escaped = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', escaped)
    escaped = re.sub(r'\*(.*?)\*', r'<i>\1</i>', escaped)
    escaped = re.sub(r'`(.*?)`', f'<font name="{FONT_MONO}" size="8.5" color="#991111"><b>\\1</b></font>', escaped)
    escaped = escaped.replace('\n', '<br/>')
    return escaped

# Helper to create styled XML elements for Word tables
def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

class NumberedCanvas(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        super(NumberedCanvas, self).__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_number(num_pages)
            super(NumberedCanvas, self).showPage()
        super(NumberedCanvas, self).save()

    def draw_page_number(self, page_count):
        if self._pageNumber == 1:
            return  # Suppress page number on cover page
        self.saveState()
        self.setFont(FONT_NORMAL, 9)
        self.setFillColor(colors.HexColor("#64748b"))
        
        # Header
        self.drawString(54, 750, "NetOracle: Closed-Loop Federated Causal Intelligence Report")
        self.setStrokeColor(colors.HexColor("#e2e8f0"))
        self.setLineWidth(0.5)
        self.line(54, 742, 558, 742)
        
        # Footer
        self.line(54, 55, 558, 55)
        page_str = f"Page {self._pageNumber} of {page_count}"
        self.drawRightString(558, 40, page_str)
        self.drawString(54, 40, "Confidential - RVCE IV Sem EL Project")
        self.restoreState()

# Markdown block parser
def parse_markdown_to_blocks(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    blocks = []
    current_block = None
    state = 'NORMAL' # NORMAL, CODE, TABLE
    
    for line in lines:
        stripped = line.strip()
        
        # Check for code blocks
        if stripped.startswith('```'):
            if state == 'CODE':
                blocks.append(current_block)
                current_block = None
                state = 'NORMAL'
            else:
                if current_block:
                    blocks.append(current_block)
                current_block = {'type': 'code', 'lines': []}
                state = 'CODE'
            continue
            
        if state == 'CODE':
            current_block['lines'].append(line.rstrip('\n'))
            continue
            
        # Check for table
        if stripped.startswith('|'):
            if state == 'TABLE':
                current_block['lines'].append(stripped)
            else:
                if current_block:
                    blocks.append(current_block)
                current_block = {'type': 'table', 'lines': [stripped]}
                state = 'TABLE'
            continue
        elif state == 'TABLE':
            blocks.append(current_block)
            current_block = None
            state = 'NORMAL'
            
        # Normal parsing
        if stripped == '':
            if current_block:
                blocks.append(current_block)
                current_block = None
            continue
            
        # Check for H1, H2, H3, H4
        if stripped.startswith('# '):
            if current_block:
                blocks.append(current_block)
            blocks.append({'type': 'h1', 'text': stripped[2:]})
            current_block = None
        elif stripped.startswith('## '):
            if current_block:
                blocks.append(current_block)
            blocks.append({'type': 'h2', 'text': stripped[3:]})
            current_block = None
        elif stripped.startswith('### '):
            if current_block:
                blocks.append(current_block)
            blocks.append({'type': 'h3', 'text': stripped[4:]})
            current_block = None
        elif stripped.startswith('#### '):
            if current_block:
                blocks.append(current_block)
            blocks.append({'type': 'h4', 'text': stripped[5:]})
            current_block = None
        # Check for blockquote
        elif stripped.startswith('> '):
            if current_block:
                if current_block['type'] == 'blockquote':
                    current_block['text'] += '\n' + stripped[2:]
                else:
                    blocks.append(current_block)
                    current_block = {'type': 'blockquote', 'text': stripped[2:]}
            else:
                current_block = {'type': 'blockquote', 'text': stripped[2:]}
        elif stripped.startswith('>'):
            if current_block:
                if current_block['type'] == 'blockquote':
                    current_block['text'] += '\n' + stripped[1:]
                else:
                    blocks.append(current_block)
                    current_block = {'type': 'blockquote', 'text': stripped[1:]}
            else:
                current_block = {'type': 'blockquote', 'text': stripped[1:]}
        # Check for list items
        elif stripped.startswith('- ') or stripped.startswith('* ') or re.match(r'^\d+\.\s', stripped):
            if current_block:
                blocks.append(current_block)
            
            main_text = stripped
            if stripped.startswith('- '):
                main_text = stripped[2:]
            elif stripped.startswith('* '):
                main_text = stripped[2:]
            else:
                match = re.match(r'^(\d+\.\s)(.*)', stripped)
                if match:
                    main_text = f"{match.group(1)}{match.group(2)}"
            
            blocks.append({'type': 'list_item', 'text': main_text})
            current_block = None
        elif stripped == '---':
            if current_block:
                blocks.append(current_block)
            blocks.append({'type': 'hr'})
            current_block = None
        else:
            if current_block:
                if current_block['type'] == 'paragraph':
                    current_block['text'] += ' ' + stripped
                else:
                    blocks.append(current_block)
                    current_block = {'type': 'paragraph', 'text': stripped}
            else:
                current_block = {'type': 'paragraph', 'text': stripped}
                
    if current_block:
        blocks.append(current_block)
        
    return blocks

def clean_table_rows(table_lines):
    parsed_rows = []
    for line in table_lines:
        parts = [p.strip() for p in line.split('|')[1:-1]]
        if all(re.match(r'^[-:\s]+$', p) for p in parts if p):
            continue
        parsed_rows.append(parts)
    return parsed_rows

# docx formatting helpers
def add_markdown_paragraph(doc, text, list_bullet=False):
    if list_bullet:
        p = doc.add_paragraph(style='List Bullet')
    else:
        p = doc.add_paragraph(style='Normal')
        
    p.paragraph_format.space_after = Pt(4 if list_bullet else 6)
    p.paragraph_format.line_spacing = 1.15
    
    # Inline formatting parser
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)')
    parts = pattern.split(text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x99, 0x11, 0x11)
        else:
            if part:
                p.add_run(part)
    return p

def add_blockquote(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, "F8FAFC")
    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
    
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="18" w:space="0" w:color="3B82F6"/>'
        f'<w:top w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'<w:bottom w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x99, 0x11, 0x11)
        else:
            if part:
                p.add_run(part)

def add_code_block(doc, code_lines):
    code_text = "\n".join(code_lines)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, "F1F5F9")
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="6" w:space="0" w:color="CBD5E1"/>'
        f'<w:top w:val="single" w:sz="6" w:space="0" w:color="CBD5E1"/>'
        f'<w:right w:val="single" w:sz="6" w:space="0" w:color="CBD5E1"/>'
        f'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="CBD5E1"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)

def add_table_block(doc, table_rows):
    if not table_rows:
        return
    
    num_cols = len(table_rows[0])
    table = doc.add_table(rows=len(table_rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    for col_idx, col_text in enumerate(table_rows[0]):
        hdr_cells[col_idx].text = col_text
        set_cell_background(hdr_cells[col_idx], "0F172A")
        set_cell_margins(hdr_cells[col_idx], top=100, bottom=100, left=120, right=120)
        p = hdr_cells[col_idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if p.runs:
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            p.runs[0].font.size = Pt(9.5)
            
    for row_idx, row_data in enumerate(table_rows[1:], start=1):
        row_cells = table.rows[row_idx].cells
        bg_color = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < len(row_cells):
                p = row_cells[col_idx].paragraphs[0]
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.15
                
                pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)')
                parts = pattern.split(cell_text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part.startswith('*') and part.endswith('*'):
                        run = p.add_run(part[1:-1])
                        run.italic = True
                    elif part.startswith('`') and part.endswith('`'):
                        run = p.add_run(part[1:-1])
                        run.font.name = 'Consolas'
                        run.font.size = Pt(8.5)
                        run.font.color.rgb = RGBColor(0x99, 0x11, 0x11)
                    else:
                        if part:
                            p.add_run(part)
                set_cell_background(row_cells[col_idx], bg_color)
                set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)

def add_title_docx(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)

def add_subtitle_docx(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(40)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

def add_footer_to_docx(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        run_left = p.add_run("NetOracle Complete Technical Report | ")
        run_left.font.name = 'Arial'
        run_left.font.size = Pt(8.5)
        run_left.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
        
        run_page = p.add_run("Page ")
        run_page.font.name = 'Arial'
        run_page.font.size = Pt(8.5)
        run_page.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
        
        fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
        instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls('w'))
        fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
        fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
        
        p._p.append(fldChar1)
        p._p.append(instrText)
        p._p.append(fldChar2)
        p._p.append(fldChar3)

# Build Word report
def build_docx_report(blocks, output_path):
    print("Building DOCX report...")
    doc = Document()
    
    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)
    
    # Cover Page
    add_title_docx(doc, "NETORACLE: COMPLETE TECHNICAL REPORT")
    add_subtitle_docx(doc, "Federated Causal 5G Network Fault Intelligence System")
    
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_before = Pt(120)
    run_meta = p_meta.add_run(
        "Course: EL IV Semester (CS-AIML)\n"
        "Institution: RV College of Engineering (RVCE)\n"
        "Author: AI Pair-Programming Assistant (Antigravity)\n"
        "Purpose: Placement Preparation & Comprehensive Project Guide\n"
        "Date: June 2026\n"
        "Document Version: 4.0 (Complete Release)\n"
    )
    run_meta.font.name = 'Arial'
    run_meta.font.size = Pt(10.5)
    run_meta.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    
    doc.add_page_break()
    
    # Render blocks
    for block in blocks:
        btype = block['type']
        
        if btype == 'h1':
            p = doc.add_paragraph()
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(block['text'])
            run.font.name = 'Arial'
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
            
        elif btype == 'h2':
            p = doc.add_paragraph()
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(block['text'])
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)
            
        elif btype == 'h3':
            p = doc.add_paragraph()
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(block['text'])
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x0f, 0x76, 0x6e)
            
        elif btype == 'h4':
            p = doc.add_paragraph()
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(block['text'])
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
            
        elif btype == 'paragraph':
            add_markdown_paragraph(doc, block['text'])
            
        elif btype == 'list_item':
            add_markdown_paragraph(doc, block['text'], list_bullet=True)
            
        elif btype == 'blockquote':
            add_blockquote(doc, block['text'])
            
        elif btype == 'code':
            add_code_block(doc, block['lines'])
            
        elif btype == 'table':
            rows = clean_table_rows(block['lines'])
            add_table_block(doc, rows)
            
        elif btype == 'hr':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run("❖   ❖   ❖")
            run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
            
    # Add page numbers in footer
    add_footer_to_docx(doc)
    doc.save(str(output_path))
    print(f"DOCX saved to {output_path}")

# Build PDF report
def build_pdf_report(blocks, output_path):
    print("Building PDF report...")
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        rightMargin=54,
        leftMargin=54,
        topMargin=54,
        bottomMargin=54
    )
    
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'CoverTitle',
        parent=styles['Normal'],
        fontName=FONT_BOLD,
        fontSize=24,
        leading=28,
        textColor=colors.HexColor("#0f172a"),
        alignment=1,
        spaceAfter=12
    )
    
    subtitle_style = ParagraphStyle(
        'CoverSubtitle',
        parent=styles['Normal'],
        fontName=FONT_ITALIC,
        fontSize=13,
        leading=16,
        textColor=colors.HexColor("#475569"),
        alignment=1,
        spaceAfter=40
    )
    
    meta_style = ParagraphStyle(
        'CoverMeta',
        parent=styles['Normal'],
        fontName=FONT_NORMAL,
        fontSize=10,
        leading=14,
        textColor=colors.HexColor("#64748b"),
        alignment=1,
        spaceAfter=120
    )

    h1_style = ParagraphStyle(
        'Heading1_Custom',
        parent=styles['Heading1'],
        fontName=FONT_BOLD,
        fontSize=15,
        leading=19,
        textColor=colors.HexColor("#0f172a"),
        spaceBefore=14,
        spaceAfter=8,
        keepWithNext=True
    )

    h2_style = ParagraphStyle(
        'Heading2_Custom',
        parent=styles['Heading2'],
        fontName=FONT_BOLD,
        fontSize=11.5,
        leading=15,
        textColor=colors.HexColor("#1e3a8a"),
        spaceBefore=12,
        spaceAfter=6,
        keepWithNext=True
    )

    h3_style = ParagraphStyle(
        'Heading3_Custom',
        parent=styles['Heading3'],
        fontName=FONT_ITALIC,
        fontSize=10,
        leading=13,
        textColor=colors.HexColor("#0f766e"),
        spaceBefore=10,
        spaceAfter=4,
        keepWithNext=True
    )
    
    h4_style = ParagraphStyle(
        'Heading4_Custom',
        parent=styles['Heading3'],
        fontName=FONT_BOLD,
        fontSize=9.5,
        leading=12,
        textColor=colors.HexColor("#475569"),
        spaceBefore=8,
        spaceAfter=2,
        keepWithNext=True
    )

    body_style = ParagraphStyle(
        'Body_Custom',
        parent=styles['Normal'],
        fontName=FONT_NORMAL,
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor("#334155"),
        spaceAfter=6
    )

    bullet_style = ParagraphStyle(
        'Bullet_Custom',
        parent=body_style,
        leftIndent=15,
        firstLineIndent=-10,
        spaceAfter=3
    )

    callout_style = ParagraphStyle(
        'Callout_Custom',
        parent=styles['Normal'],
        fontName=FONT_ITALIC,
        fontSize=9,
        leading=12.5,
        textColor=colors.HexColor("#1e293b")
    )
    
    code_style = ParagraphStyle(
        'Code_Custom',
        parent=styles['Normal'],
        fontName=FONT_MONO,
        fontSize=8,
        leading=11,
        textColor=colors.HexColor("#1e293b")
    )

    story = []

    # Cover Page
    story.append(Spacer(1, 120))
    story.append(Paragraph("NETORACLE: COMPLETE TECHNICAL REPORT", title_style))
    story.append(Paragraph("Federated Causal 5G Network Fault Intelligence System", subtitle_style))
    story.append(Spacer(1, 40))
    
    meta_text = (
        "<b>Course:</b> EL IV Semester (CS-AIML)<br/>"
        "<b>Institution:</b> RV College of Engineering (RVCE)<br/>"
        "<b>Author:</b> AI Pair-Programming Assistant (Antigravity)<br/>"
        "<b>Purpose:</b> Placement Preparation & Comprehensive Project Guide<br/>"
        "<b>Date:</b> June 2026<br/>"
        "<b>Version:</b> 4.0 (Complete Release)<br/>"
    )
    story.append(Paragraph(meta_text, meta_style))
    story.append(PageBreak())

    # Render blocks
    for block in blocks:
        btype = block['type']
        
        if btype == 'h1':
            story.append(Paragraph(html_escape_and_md_to_html(block['text']), h1_style))
            
        elif btype == 'h2':
            story.append(Paragraph(html_escape_and_md_to_html(block['text']), h2_style))
            
        elif btype == 'h3':
            story.append(Paragraph(html_escape_and_md_to_html(block['text']), h3_style))
            
        elif btype == 'h4':
            story.append(Paragraph(html_escape_and_md_to_html(block['text']), h4_style))
            
        elif btype == 'paragraph':
            text = html_escape_and_md_to_html(block['text'])
            story.append(Paragraph(text, body_style))
            
        elif btype == 'list_item':
            text = html_escape_and_md_to_html(block['text'])
            story.append(Paragraph(f"• {text}", bullet_style))
            
        elif btype == 'blockquote':
            text = html_escape_and_md_to_html(block['text'])
            callout_data = [[Paragraph(text, callout_style)]]
            callout_table = Table(callout_data, colWidths=[490])
            callout_table.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,-1), colors.HexColor("#f8fafc")),
                ('LEFTPADDING', (0,0), (-1,-1), 12),
                ('RIGHTPADDING', (0,0), (-1,-1), 12),
                ('TOPPADDING', (0,0), (-1,-1), 8),
                ('BOTTOMPADDING', (0,0), (-1,-1), 8),
                ('LINELEFT', (0,0), (0,-1), 3, colors.HexColor("#3b82f6")),
            ]))
            story.append(Spacer(1, 4))
            story.append(callout_table)
            story.append(Spacer(1, 4))
            
        elif btype == 'code':
            code_text = html.escape("\n".join(block['lines']))
            p_code = Preformatted(code_text, code_style)
            code_table = Table([[p_code]], colWidths=[490])
            code_table.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,-1), colors.HexColor("#f1f5f9")),
                ('LEFTPADDING', (0,0), (-1,-1), 10),
                ('RIGHTPADDING', (0,0), (-1,-1), 10),
                ('TOPPADDING', (0,0), (-1,-1), 8),
                ('BOTTOMPADDING', (0,0), (-1,-1), 8),
                ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor("#cbd5e1")),
            ]))
            story.append(Spacer(1, 4))
            story.append(code_table)
            story.append(Spacer(1, 4))
            
        elif btype == 'table':
            rows = clean_table_rows(block['lines'])
            table_data = []
            for row_idx, row in enumerate(rows):
                table_row = []
                for cell in row:
                    text = html_escape_and_md_to_html(cell)
                    if row_idx == 0:
                        table_row.append(Paragraph(f"<b>{text}</b>", ParagraphStyle('Hdr', parent=body_style, textColor=colors.white, fontName=FONT_BOLD)))
                    else:
                        table_row.append(Paragraph(text, body_style))
                table_data.append(table_row)
                
            if table_data:
                num_cols = len(table_data[0])
                col_width = 490 / num_cols
                col_widths = [col_width] * num_cols
                
                pdf_table = Table(table_data, colWidths=col_widths)
                t_style = [
                    ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#0f172a")),
                    ('VALIGN', (0,0), (-1,-1), 'TOP'),
                    ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#cbd5e1")),
                    ('TOPPADDING', (0,0), (-1,-1), 6),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 6),
                    ('LEFTPADDING', (0,0), (-1,-1), 8),
                    ('RIGHTPADDING', (0,0), (-1,-1), 8),
                ]
                for r in range(1, len(table_data)):
                    if r % 2 == 1:
                        t_style.append(('BACKGROUND', (0,r), (-1,r), colors.HexColor("#f8fafc")))
                pdf_table.setStyle(TableStyle(t_style))
                story.append(Spacer(1, 4))
                story.append(pdf_table)
                story.append(Spacer(1, 4))
                
        elif btype == 'hr':
            hr_table = Table([['']], colWidths=[200])
            hr_table.setStyle(TableStyle([
                ('LINEABOVE', (0,0), (-1,-1), 1, colors.HexColor("#cbd5e1")),
                ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                ('BOTTOMPADDING', (0,0), (-1,-1), 0),
                ('TOPPADDING', (0,0), (-1,-1), 0),
            ]))
            story.append(Spacer(1, 10))
            story.append(hr_table)
            story.append(Spacer(1, 10))

    doc.build(story, canvasmaker=NumberedCanvas)
    print(f"PDF saved to {output_path}")

def main():
    if not MD_PATH.exists():
        print(f"Error: Markdown report file not found at {MD_PATH}")
        sys.exit(1)
        
    print(f"Reading markdown from {MD_PATH}...")
    blocks = parse_markdown_to_blocks(MD_PATH)
    print(f"Parsed {len(blocks)} blocks.")
    
    # Build documents in reports/ directory
    build_docx_report(blocks, DOCX_OUT_PATH)
    build_pdf_report(blocks, PDF_OUT_PATH)
    
    # Copy documents to artifact/brain directory
    print(f"Copying files to artifact directory...")
    shutil.copy2(DOCX_OUT_PATH, DOCX_ART_PATH)
    shutil.copy2(PDF_OUT_PATH, PDF_ART_PATH)
    print("Files successfully copied to artifact directory.")

if __name__ == "__main__":
    main()
